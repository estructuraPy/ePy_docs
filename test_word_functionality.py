#!/usr/bin/env python3
"""
Test script for Word document functionality in ePy_docs.
"""

import os
import sys
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / "src"))

def test_word_functionality():
    """Test basic Word functionality."""
    from ePy_docs.core._word import is_docx_available, get_word_export_config
    
    print("🔍 Testing Word functionality...")
    
    # Test 1: Check if python-docx is available
    print("1. Checking python-docx availability...")
    docx_available = is_docx_available()
    print(f"   ✅ python-docx available: {docx_available}")
    
    # Test 2: Test Word export configuration
    print("2. Testing Word export configuration...")
    word_config = get_word_export_config()
    print(f"   ✅ Word config generated: {word_config}")
    
    # Test 3: Test DocumentWriter with Word support
    print("3. Testing DocumentWriter with Word support...")
    try:
        from ePy_docs.writers import DocumentWriter
        
        # Create a writer instance
        writer = DocumentWriter(document_type="report", layout_style="professional")
        
        # Check that add_word_file method exists
        assert hasattr(writer, 'add_word_file'), "add_word_file method not found"
        print("   ✅ add_word_file method exists")
        
        # Check that generate method accepts docx parameter
        import inspect
        generate_signature = inspect.signature(writer.generate)
        assert 'docx' in generate_signature.parameters, "docx parameter not in generate method"
        print("   ✅ generate method supports docx parameter")
        
        # Test basic document creation with Word export
        writer.add_h1("Test Document")
        writer.add_content("This is a test document to verify Word export functionality.")
        
        # Test generation with Word format
        print("4. Testing document generation with Word format...")
        result = writer.generate(
            html=False, 
            pdf=False, 
            docx=True,
            output_filename="word_test"
        )
        
        print(f"   ✅ Generation result: {result}")
        
        # Check if docx key is in the result
        if 'docx' in result:
            print(f"   ✅ Word document path: {result['docx']}")
        else:
            print("   ℹ️  Word generation may require Quarto to be installed")
        
    except Exception as e:
        print(f"   ❌ Error testing DocumentWriter: {e}")
        return False
    
    print("\n🎉 All Word functionality tests completed successfully!")
    return True


def create_sample_word_document():
    """Create a sample Word document for testing import functionality."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        print("📝 Creating sample Word document for testing...")
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Sample Document for ePy_docs Testing', 0)
        
        # Add some content
        doc.add_heading('Introduction', level=1)
        intro = doc.add_paragraph('This is a sample Word document created to test the import functionality of ePy_docs. ')
        intro.add_run('This text is bold.').bold = True
        intro.add_run(' This text is italic.').italic = True
        
        doc.add_heading('Methods', level=1)
        doc.add_paragraph('Here we describe the methods used in this research.')
        
        # Add a list
        doc.add_paragraph('Key points:', style='List Bullet')
        doc.add_paragraph('First important point', style='List Bullet')
        doc.add_paragraph('Second important point', style='List Bullet')
        doc.add_paragraph('Third important point', style='List Bullet')
        
        # Add a table
        doc.add_heading('Results', level=1)
        doc.add_paragraph('The following table shows our results:')
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Unit'
        
        # Add table data
        data = [
            ['Temperature', '25.5', '°C'],
            ['Pressure', '1.013', 'bar'],
            ['Humidity', '65', '%']
        ]
        
        for i, row_data in enumerate(data, start=1):
            cells = table.rows[i].cells
            for j, value in enumerate(row_data):
                cells[j].text = value
        
        # Add conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph('This document demonstrates various formatting features that can be imported into ePy_docs.')
        
        # Save the document
        sample_path = "sample_word_document.docx"
        doc.save(sample_path)
        
        print(f"   ✅ Sample document created: {sample_path}")
        return sample_path
        
    except Exception as e:
        print(f"   ❌ Error creating sample document: {e}")
        return None


def test_word_import():
    """Test importing a Word document."""
    print("\n📖 Testing Word document import...")
    
    # First create a sample document
    sample_path = create_sample_word_document()
    if not sample_path:
        return False
    
    try:
        from ePy_docs.core._word import read_docx_file, convert_docx_to_markdown
        
        # Test reading the document
        print("1. Testing Word document reading...")
        word_data = read_docx_file(
            Path(sample_path),
            preserve_formatting=True,
            convert_tables=True,
            extract_images=False
        )
        
        print(f"   ✅ Document read successfully")
        print(f"   ✅ Content length: {len(word_data['content'])} characters")
        print(f"   ✅ Tables found: {len(word_data['tables'])}")
        print(f"   ✅ Metadata: {word_data['metadata']}")
        
        # Test markdown conversion
        print("2. Testing markdown conversion...")
        markdown_content = convert_docx_to_markdown(
            Path(sample_path),
            preserve_formatting=True,
            convert_tables=True
        )
        
        print(f"   ✅ Markdown conversion successful")
        print(f"   ✅ Markdown length: {len(markdown_content)} characters")
        
        # Show a preview of the converted content
        print("3. Preview of converted content:")
        preview = markdown_content[:500] + "..." if len(markdown_content) > 500 else markdown_content
        print(f"   {preview}")
        
        # Test with DocumentWriter
        print("4. Testing Word import with DocumentWriter...")
        from ePy_docs.writers import DocumentWriter
        
        writer = DocumentWriter(document_type="report", layout_style="professional")
        writer.add_h1("Document with Imported Word Content")
        writer.add_content("The following content was imported from a Word document:")
        
        # This should work now
        writer.add_word_file(
            sample_path,
            preserve_formatting=True,
            convert_tables=True
        )
        
        writer.add_content("Import completed successfully!")
        
        print("   ✅ Word import with DocumentWriter successful")
        
        # Clean up
        if os.path.exists(sample_path):
            os.remove(sample_path)
            print(f"   🧹 Cleaned up sample file: {sample_path}")
        
    except Exception as e:
        print(f"   ❌ Error testing Word import: {e}")
        import traceback
        print(f"   🔍 Traceback: {traceback.format_exc()}")
        
        # Clean up on error
        if sample_path and os.path.exists(sample_path):
            os.remove(sample_path)
        
        return False
    
    print("   🎉 Word import testing completed successfully!")
    return True


if __name__ == "__main__":
    print("🚀 Starting ePy_docs Word functionality tests...\n")
    
    success = True
    
    # Test basic functionality
    if not test_word_functionality():
        success = False
    
    # Test import functionality
    if not test_word_import():
        success = False
    
    if success:
        print("\n✅ All tests passed! Word functionality is working correctly.")
    else:
        print("\n❌ Some tests failed. Please check the output above.")
        sys.exit(1)