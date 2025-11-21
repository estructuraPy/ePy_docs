"""
Word Document Processing Module

Handles:
- DOCX file reading and content extraction
- Advanced conversion to markdown/QMD format with formatting preservation
- Table extraction and conversion
- Image extraction and processing
- Word document export support via Quarto
"""

import os
import re
import warnings
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path


# =============================================================================
# WORD DOCUMENT READING
# =============================================================================

def read_docx_file(file_path: Path, preserve_formatting: bool = True, 
                   convert_tables: bool = True, extract_images: bool = False,
                   image_output_dir: Optional[Path] = None) -> Dict[str, Any]:
    """
    Read and extract comprehensive content from Word document (.docx).
    
    Args:
        file_path: Path to DOCX file
        preserve_formatting: If True, preserves basic formatting (bold, italic)
        convert_tables: If True, converts Word tables to Markdown format
        extract_images: If True, extracts embedded images
        image_output_dir: Directory to save extracted images (required if extract_images=True)
        
    Returns:
        Dictionary containing:
        - 'content': Markdown-formatted text content
        - 'tables': List of extracted tables (if convert_tables=True)
        - 'images': List of extracted image paths (if extract_images=True)
        - 'metadata': Document metadata
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ImportError: If python-docx not installed
        ValueError: If extract_images=True but image_output_dir not provided
    """
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    if extract_images and image_output_dir is None:
        raise ValueError("image_output_dir required when extract_images=True")
    
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx library required for Word document processing.\n"
            "Install with: pip install python-docx"
        )
    
    doc = Document(file_path)
    
    # Initialize result with ordered elements
    result = {
        'content': '',
        'tables': [],
        'images': [],
        'elements': [],  # Ordered list of elements with their types
        'metadata': _extract_document_metadata(doc)
    }
    
    content_parts = []
    element_index = 0
    
    # Process document elements in order to maintain structure
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            paragraph = _find_paragraph_by_element(doc, element)
            if paragraph:
                # Check if paragraph contains images
                has_images = False
                if extract_images:
                    images_in_para = _extract_images_from_paragraph(paragraph, image_output_dir, element_index, doc)
                    if images_in_para:
                        has_images = True
                        for img_info in images_in_para:
                            result['images'].append(img_info)
                            result['elements'].append({
                                'type': 'image',
                                'index': element_index,
                                'data': img_info
                            })
                            element_index += 1
                
                # Extract text content (skip if paragraph only contains image)
                para_text = _extract_paragraph_with_formatting(paragraph, preserve_formatting)
                if para_text.strip() and not (has_images and len(para_text.strip()) < 10):
                    # Handle headings
                    if paragraph.style.name.startswith('Heading'):
                        level = _extract_heading_level(paragraph.style.name)
                        formatted_text = f"{'#' * level} {para_text}"
                    else:
                        formatted_text = para_text
                    
                    content_parts.append(formatted_text)
                    result['elements'].append({
                        'type': 'text',
                        'index': element_index,
                        'data': formatted_text
                    })
                    element_index += 1
        
        elif element.tag.endswith('tbl'):  # Table
            table = _find_table_by_element(doc, element)
            if table:
                # Check if this is actually a callout disguised as a table
                if _is_callout_table(table):
                    # Process as callout
                    callout_md = _convert_callout_to_markdown(table)
                    if callout_md:
                        content_parts.append(callout_md)
                        result['elements'].append({
                            'type': 'callout',
                            'index': element_index,
                            'data': callout_md
                        })
                        element_index += 1
                else:
                    # Process as regular table
                    table_data = _extract_table_data(table)
                    if table_data and len(table_data) > 0:
                        result['tables'].append(table_data)
                        result['elements'].append({
                            'type': 'table',
                            'index': element_index,
                            'data': table_data
                        })
                        element_index += 1
                        
                        # Only add markdown if convert_tables is False
                        if convert_tables:
                            table_md, _ = _convert_table_to_markdown(table)
                            if table_md:
                                content_parts.append(table_md)
    
    result['content'] = '\n\n'.join(content_parts)
    return result


# =============================================================================
# CONTENT EXTRACTION
# =============================================================================

def extract_paragraphs(file_path: Path) -> List[str]:
    """
    Extract paragraphs from Word document.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        List of paragraph texts
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library required")
    
    doc = Document(file_path)
    return [p.text for p in doc.paragraphs if p.text.strip()]


def extract_tables(file_path: Path) -> List[List[List[str]]]:
    """
    Extract tables from Word document.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        List of tables, each table is a list of rows, each row is a list of cell values
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library required")
    
    doc = Document(file_path)
    tables = []
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    return tables


def extract_headings(file_path: Path) -> List[Dict[str, Any]]:
    """
    Extract headings from Word document.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        List of headings with {'level': int, 'text': str}
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library required")
    
    doc = Document(file_path)
    headings = []
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.replace('Heading ', ''))
                headings.append({
                    'level': level,
                    'text': paragraph.text
                })
            except ValueError:
                continue
    
    return headings


# =============================================================================
# CONVERSION TO MARKDOWN
# =============================================================================

def convert_docx_to_markdown(file_path: Path, preserve_formatting: bool = True,
                            convert_tables: bool = True) -> str:
    """
    Convert Word document to markdown format with enhanced processing.
    
    Args:
        file_path: Path to DOCX file
        preserve_formatting: If True, preserves basic formatting (bold, italic)
        convert_tables: If True, converts Word tables to Markdown format
        
    Returns:
        Markdown content
    """
    # Use the enhanced read function
    result = read_docx_file(
        file_path, 
        preserve_formatting=preserve_formatting,
        convert_tables=convert_tables,
        extract_images=False
    )
    
    # Process and clean up the content
    content = process_word_content(result['content'], fix_formatting=True)
    
    return content


def convert_docx_to_qmd(
    docx_path: Path,
    output_path: Path,
    yaml_config: Optional[Dict[str, Any]] = None
) -> Path:
    """
    Convert Word document to Quarto QMD file.
    
    Args:
        docx_path: Path to input DOCX file
        output_path: Path to output QMD file
        yaml_config: Optional YAML configuration
        
    Returns:
        Path to created QMD file
    """
    import yaml
    
    # Convert to markdown
    markdown_content = convert_docx_to_markdown(docx_path)
    
    # Add YAML frontmatter
    if yaml_config is None:
        yaml_config = {
            'title': docx_path.stem.replace('_', ' ').title()
        }
    
    yaml_str = yaml.dump(yaml_config, default_flow_style=False, sort_keys=False)
    
    qmd_content = f'''---
{yaml_str}---

{markdown_content}
'''
    
    # Write to file
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(qmd_content)
    
    return output_path


# =============================================================================
# UTILITIES
# =============================================================================

def validate_docx_file(file_path: Path) -> bool:
    """
    Validate that file is a Word document.
    
    Args:
        file_path: Path to file
        
    Returns:
        True if valid DOCX file
        
    Raises:
        ValueError: If not a DOCX file
    """
    if file_path.suffix.lower() != '.docx':
        raise ValueError(
            f"Not a Word document: {file_path}\n"
            f"Expected .docx file"
        )
    
    return True


def is_docx_available() -> bool:
    """
    Check if python-docx library is available.
    
    Returns:
        True if python-docx is installed
    """
    try:
        import docx
        return True
    except ImportError:
        return False


# =============================================================================
# HELPER FUNCTIONS FOR ENHANCED DOCX PROCESSING
# =============================================================================

def _extract_document_metadata(doc) -> Dict[str, Any]:
    """Extract metadata from Word document."""
    metadata = {}
    
    # Core properties
    core_props = doc.core_properties
    metadata.update({
        'title': core_props.title or '',
        'author': core_props.author or '',
        'subject': core_props.subject or '',
        'keywords': core_props.keywords or '',
        'created': core_props.created.isoformat() if core_props.created else None,
        'modified': core_props.modified.isoformat() if core_props.modified else None,
    })
    
    return metadata


def _find_paragraph_by_element(doc, element):
    """Find paragraph object by its XML element."""
    for paragraph in doc.paragraphs:
        if paragraph._element == element:
            return paragraph
    return None


def _find_table_by_element(doc, element):
    """Find table object by its XML element."""
    for table in doc.tables:
        if table._element == element:
            return table
    return None


def _extract_paragraph_with_formatting(paragraph, preserve_formatting: bool = True) -> str:
    """Extract text from paragraph with optional formatting preservation."""
    if not preserve_formatting:
        return paragraph.text
    
    text_parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        # Apply markdown formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        # Handle strikethrough (requires checking the run's element)
        try:
            if run._element.rPr is not None and run._element.rPr.strike is not None:
                text = f"~~{text}~~"
        except:
            pass  # Not all runs have these properties
        
        text_parts.append(text)
    
    return ''.join(text_parts)


def _extract_heading_level(style_name: str) -> int:
    """Extract heading level from style name."""
    match = re.search(r'Heading (\d+)', style_name)
    if match:
        return int(match.group(1))
    return 1  # Default to H1


def _convert_table_to_markdown(table) -> Tuple[str, List[List[str]]]:
    """Convert Word table to Markdown format and return both markdown and raw data."""
    if not table.rows:
        return "", []
    
    # Extract table data
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    
    if not table_data:
        return "", []
    
    # Generate markdown table
    markdown_rows = []
    for i, row_data in enumerate(table_data):
        # Handle empty cells
        cleaned_row = [cell if cell else " " for cell in row_data]
        markdown_rows.append('| ' + ' | '.join(cleaned_row) + ' |')
        
        # Add separator after header row
        if i == 0 and len(table_data) > 1:
            separator = '| ' + ' | '.join(['---'] * len(cleaned_row)) + ' |'
            markdown_rows.append(separator)
    
    return '\n'.join(markdown_rows), table_data


def _extract_images_from_docx(doc, output_dir: Path) -> List[str]:
    """Extract images from Word document and save them to output directory."""
    output_dir.mkdir(parents=True, exist_ok=True)
    image_paths = []
    
    try:
        # Extract images from document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    img_data = rel.target_part.blob
                    
                    # Determine file extension from content type
                    content_type = rel.target_part.content_type
                    if 'png' in content_type:
                        ext = '.png'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    else:
                        ext = '.png'  # Default to PNG
                    
                    img_filename = f"extracted_image_{len(image_paths) + 1}{ext}"
                    img_path = output_dir / img_filename
                    
                    with open(img_path, 'wb') as img_file:
                        img_file.write(img_data)
                    
                    image_paths.append(str(img_path))
                    
                except Exception as e:
                    warnings.warn(f"Failed to extract image: {e}")
                    continue
    
    except Exception as e:
        warnings.warn(f"Error extracting images: {e}")
    
    return image_paths


def _extract_images_from_paragraph(paragraph, output_dir: Path, base_index: int = 0, doc=None) -> List[Dict[str, Any]]:
    """Extract images from a specific paragraph and return image info."""
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)
    
    images = []
    
    try:
        # Look for images in runs
        for run in paragraph.runs:
            # Check if run contains drawing elements
            drawings = run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing', 
                                         run.element.nsmap if hasattr(run.element, 'nsmap') else {})
            
            if not drawings:
                # Try alternative namespace
                from docx.oxml import parse_xml
                drawings = run._element.findall('.//{*}drawing')
            for drawing in drawings:
                try:
                    # Extract embedded image reference
                    blips = drawing.findall('.//{*}blip')
                    for blip in blips:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed:
                            # Get image from relationship using document part
                            if doc and hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
                                image_part = doc.part.related_parts[embed]
                                img_data = image_part.blob
                            else:
                                # Fallback: try to get from paragraph's parent
                                try:
                                    image_part = paragraph._parent._element.part.related_parts[embed]
                                    img_data = image_part.blob
                                except AttributeError:
                                    # Skip this image if we can't access the part
                                    continue
                            
                            # Determine extension
                            content_type = image_part.content_type
                            if 'png' in content_type:
                                ext = '.png'
                            elif 'jpeg' in content_type or 'jpg' in content_type:
                                ext = '.jpg'
                            elif 'gif' in content_type:
                                ext = '.gif'
                            else:
                                ext = '.png'
                            
                            # Save image
                            img_filename = f"extracted_image_{base_index}_{len(images) + 1}{ext}"
                            img_path = output_dir / img_filename if output_dir else Path(img_filename)
                            
                            if output_dir:
                                with open(img_path, 'wb') as img_file:
                                    img_file.write(img_data)
                            
                            images.append({
                                'path': str(img_path),
                                'alt_text': 'Image from Word document'
                            })
                except Exception as e:
                    warnings.warn(f"Failed to extract image from paragraph: {e}")
                    continue
    except Exception as e:
        warnings.warn(f"Error processing paragraph images: {e}")
    
    return images


def _is_callout_table(table) -> bool:
    """Determine if a Word table is actually a callout/admonition."""
    if not table.rows:
        return False
    
    # Check table dimensions - callouts are typically single column or very narrow
    max_cols = max(len(row.cells) for row in table.rows) if table.rows else 0
    if max_cols > 2:
        return False  # Real tables usually have more columns
    
    # Check content patterns - callouts often have keywords
    all_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_text += cell.text.lower() + " "
    
    # Common callout keywords in Spanish and English
    callout_keywords = [
        'nota:', 'note:', 'observación:', 'observation:',
        'importante:', 'important:', 'atención:', 'attention:',
        'advertencia:', 'warning:', 'precaución:', 'caution:',
        'tip:', 'consejo:', 'ejemplo:', 'example:',
        'info:', 'información:', 'information:',
        'recordatorio:', 'reminder:', 'aviso:', 'notice:'
    ]
    
    # Check if content starts with typical callout words
    for keyword in callout_keywords:
        if keyword in all_text:
            return True
    
    # Check table styling (simple heuristic based on structure)
    # Callouts often have 1-2 rows with merged cells or single column
    if len(table.rows) <= 2 and max_cols == 1:
        return True
    
    # Check if it's a single-cell table with significant text (likely a callout)
    if len(table.rows) == 1 and max_cols == 1:
        cell_text = table.rows[0].cells[0].text.strip()
        if len(cell_text) > 20:  # Substantial text content
            return True
    
    return False


def _convert_callout_to_markdown(table) -> str:
    """Convert a Word table that represents a callout to Quarto callout format."""
    if not table.rows:
        return ""
    
    # Extract all text from the table
    all_text = []
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                all_text.append(cell_text)
    
    content = " ".join(all_text)
    if not content:
        return ""
    
    # Determine callout type based on content
    content_lower = content.lower()
    callout_type = "note"  # default
    
    if any(word in content_lower for word in ['importante', 'important', 'atención', 'attention']):
        callout_type = "important"
    elif any(word in content_lower for word in ['advertencia', 'warning', 'precaución', 'caution']):
        callout_type = "warning"
    elif any(word in content_lower for word in ['tip', 'consejo']):
        callout_type = "tip"
    elif any(word in content_lower for word in ['ejemplo', 'example']):
        callout_type = "note"  # Use note for examples
    elif any(word in content_lower for word in ['info', 'información', 'information']):
        callout_type = "note"
    
    # Clean up the content - remove the type indicator if it's at the beginning
    cleaned_content = content
    for prefix in ['Nota:', 'Note:', 'Importante:', 'Important:', 'Advertencia:', 'Warning:', 
                   'Tip:', 'Consejo:', 'Ejemplo:', 'Example:', 'Info:', 'Información:', 'Information:']:
        if cleaned_content.startswith(prefix):
            cleaned_content = cleaned_content[len(prefix):].strip()
            break
    
    # Generate Quarto callout format
    callout_markdown = f"::: {{.callout-{callout_type}}}\n{cleaned_content}\n:::"
    
    return callout_markdown


def _extract_table_data(table) -> List[List[str]]:
    """Extract raw data from a Word table."""
    if not table.rows:
        return []
    
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    
    return table_data


# =============================================================================
# CONTENT PROCESSING AND CLEANUP
# =============================================================================

def process_word_content(content: str, fix_formatting: bool = True) -> str:
    """Process extracted Word content to fix common conversion issues.
    
    Args:
        content: Markdown content extracted from Word
        fix_formatting: If True, applies formatting fixes
        
    Returns:
        Processed markdown content
    """
    if not fix_formatting:
        return content
    
    processed = content
    
    # Remove excessive whitespace
    processed = re.sub(r'\n{3,}', '\n\n', processed)
    processed = re.sub(r'[ \t]+\n', '\n', processed)  # Remove trailing spaces
    
    # Fix common bullet point issues
    processed = re.sub(r'^\s*•\s+', '- ', processed, flags=re.MULTILINE)
    processed = re.sub(r'^\s*◦\s+', '  - ', processed, flags=re.MULTILINE)
    processed = re.sub(r'^\s*▪\s+', '    - ', processed, flags=re.MULTILINE)
    
    # Fix numbered lists
    processed = re.sub(r'^\s*(\d+)\.\s+', r'\1. ', processed, flags=re.MULTILINE)
    
    # Fix common quotation mark issues
    processed = processed.replace('"', '"').replace('"', '"')
    processed = processed.replace(''', "'").replace(''', "'")
    
    # Fix em dashes and en dashes
    processed = processed.replace('—', ' -- ').replace('–', ' - ')
    
    return processed.strip()


# =============================================================================
# WORD EXPORT SUPPORT VIA QUARTO
# =============================================================================

def get_docx_config(layout_name: Optional[str] = None, reference_doc: Optional[str] = None, **kwargs) -> Dict[str, Any]:
    """Get Word (docx) configuration from layout settings.
    
    Configuration is merged from:
    1. quarto_common (toc, number-sections) from layout
    2. quarto_docx (format-specific overrides) from layout
    3. reference_doc parameter
    4. kwargs overrides
    
    Args:
        layout_name: Layout style name (e.g., 'academic', 'professional', 'classic')
        reference_doc: Path to Word reference document template (optional)
        **kwargs: Additional configuration options
        
    Returns:
        Dictionary with Word export configuration for Quarto
        
    Raises:
        ValueError: If layout or required configuration sections are missing
    """
    from ePy_docs.core._config import get_layout_config, get_config_section
    
    # Determine layout_name if not provided
    if layout_name is None:
        layout_name = 'classic'  # Default layout
    
    # Get layout configuration
    layout_config = get_layout_config(layout_name)
    
    # Validate required sections
    if 'quarto_common' not in layout_config:
        raise ValueError(
            f"Missing 'quarto_common' section in layout '{layout_name}'. "
            f"Please add quarto_common configuration to layouts/{layout_name}.epyson"
        )
    
    if 'quarto_docx' not in layout_config:
        raise ValueError(
            f"Missing 'quarto_docx' section in layout '{layout_name}'. "
            f"Please add quarto_docx configuration to layouts/{layout_name}.epyson"
        )
    
    # Merge: quarto_common → quarto_docx → reference_doc → kwargs
    merged_config = layout_config['quarto_common'].copy()
    merged_config.update(layout_config['quarto_docx'])
    
    # Add reference-doc if provided
    if reference_doc:
        merged_config['reference-doc'] = reference_doc
    
    # Override with any kwargs
    merged_config.update(kwargs)
    
    return merged_config


def add_word_format_to_yaml(yaml_config: Dict[str, Any], 
                           layout_name: Optional[str] = None,
                           reference_doc: Optional[str] = None) -> Dict[str, Any]:
    """Add Word format configuration to existing YAML config.
    
    Args:
        yaml_config: Existing YAML configuration
        layout_name: Layout style name (if None, uses default_layout)
        reference_doc: Path to Word reference document template
        
    Returns:
        Updated YAML configuration with Word format
        
    Raises:
        ValueError: If docx configuration is missing
    """
    if 'format' not in yaml_config:
        yaml_config['format'] = {}
    
    # Get word configuration from documents.epyson
    word_config = get_docx_config(layout_name=layout_name, reference_doc=reference_doc)
    
    yaml_config['format']['docx'] = word_config
    
    return yaml_config


# =============================================================================
# INTEGRATION WITH ePy_docs WORKFLOW
# =============================================================================

def process_word_file(file_path: str, preserve_formatting: bool = True,
                     convert_tables: bool = True, extract_images: bool = True,
                     image_output_dir: Optional[str] = None, fix_image_paths: bool = True,
                     output_dir: Path = None, writer_instance=None, show_figure: bool = False):
    """Process a Word file and integrate it into the ePy_docs workflow.
    
    This function is called by the DocumentWriter.add_word_file() method.
    
    Args:
        file_path: Path to the Word document (.docx)
        preserve_formatting: If True, preserves basic formatting
        convert_tables: If True, converts tables appropriately for the workflow
        extract_images: If True, extracts embedded images
        image_output_dir: Directory to save extracted images
        fix_image_paths: If True, fixes image paths for the document context
        output_dir: Output directory for the document (from writer instance)
        writer_instance: Reference to the DocumentWriter instance
    """
    file_path = Path(file_path)
    
    # Validate file exists and is a docx file
    validate_docx_file(file_path)
    
    # Check if python-docx is available
    if not is_docx_available():
        raise ImportError(
            "python-docx library required for Word document processing.\n"
            "Install with: pip install python-docx"
        )
    
    # Set up image extraction if requested
    if extract_images:
        if image_output_dir is None:
            # Use default image directory in output
            if output_dir:
                output_dir_path = Path(output_dir) if isinstance(output_dir, str) else output_dir
                image_output_dir = output_dir_path / "images" / "extracted_from_word"
            else:
                image_output_dir = Path("images") / "extracted_from_word"
        else:
            image_output_dir = Path(image_output_dir)
    
    # Extract content from Word document
    try:
        word_data = read_docx_file(
            file_path,
            preserve_formatting=preserve_formatting,
            convert_tables=False,  # Don't convert to markdown - we'll process separately
            extract_images=extract_images,
            image_output_dir=image_output_dir
        )
        
        # Add spacer before imported content
        if writer_instance and writer_instance.content_buffer:
            writer_instance.content_buffer.append('\n\n---\n\n')
        
        # Process elements in order to maintain document structure
        if writer_instance and 'elements' in word_data:
            for element in word_data['elements']:
                element_type = element['type']
                element_data = element['data']
                
                if element_type == 'text':
                    # Add text content directly
                    writer_instance.content_buffer.append(element_data + '\n\n')
                
                elif element_type == 'callout':
                    # Add callout content directly as markdown
                    writer_instance.content_buffer.append(element_data + '\n\n')
                
                elif element_type == 'table' and convert_tables:
                    # Process table with add_table
                    try:
                        import pandas as pd
                        if len(element_data) >= 2:  # Need header + at least 1 row
                            df = pd.DataFrame(element_data[1:], columns=element_data[0])
                            writer_instance.add_table(
                                df=df,
                                title=None,
                                show_figure=show_figure
                            )
                    except Exception as e:
                        warnings.warn(f"Could not process table: {e}")
                
                elif element_type == 'image' and extract_images:
                    # Process image with add_image
                    try:
                        img_path = element_data.get('path')
                        alt_text = element_data.get('alt_text', 'Image from Word document')
                        if img_path:
                            # Pass absolute path to add_image - it will handle copying and renaming
                            # No need to make it relative here, as _process_image_file expects
                            # an absolute path or path relative to current working directory
                            writer_instance.add_image(
                                path=str(img_path),
                                caption=None,
                                alt_text=alt_text,
                                responsive=True
                            )
                    except Exception as e:
                        warnings.warn(f"Could not process image: {e}")
        
        # Add final spacer
        if writer_instance:
            writer_instance.content_buffer.append('\n\n---\n\n')
            
    except Exception as e:
        raise RuntimeError(f"Failed to process Word document '{file_path}': {str(e)}")


def _process_extracted_images(image_paths: List[str], writer_instance, fix_paths: bool = True):
    """Process extracted images and add them to the document."""
    if not writer_instance or not image_paths:
        return
    
    # Add each extracted image using the writer's add_image method
    for img_path in image_paths:
        try:
            writer_instance.add_image(
                path=str(img_path),
                caption=None,  # Could be enhanced to extract captions from alt text
                width=None,    # Use default width
                alt_text="Extracted from Word document",
                responsive=True
            )
        except Exception as e:
            # If adding image fails, log it but continue with others
            print(f"Warning: Could not add extracted image {img_path}: {e}")


def _process_extracted_tables(tables: List[List[List[str]]], writer_instance, show_figure: bool = False):
    """Process extracted table data and add them to the document."""
    if not writer_instance or not tables:
        return
    
    import pandas as pd
    
    # Add each table using the writer's add_table method
    for i, table_data in enumerate(tables, 1):
        try:
            if len(table_data) < 2:  # Need at least header + 1 row
                continue
            
            # First row is header, rest are data
            df = pd.DataFrame(table_data[1:], columns=table_data[0])
            
            # Add table with automatic numbering
            writer_instance.add_table(
                df=df,
                title=None,  # Auto-generate "Table N"
                show_figure=show_figure,
                columns=None,
                max_rows_per_table=None,
                hide_columns=None,
                filter_by=None,
                sort_by=None,
                width_inches=None
            )
        except Exception as e:
            # If adding table fails, log it but continue with others
            print(f"Warning: Could not add extracted table {i}: {e}")
